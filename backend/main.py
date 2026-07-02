import pytesseract
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from PIL import Image
import io
import pdfplumber
import docx
import spacy
from spacy.matcher import PhraseMatcher
from spacy.util import filter_spans
import json
from pathlib import Path
import re
from rapidfuzz import fuzz
import os
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, ValidationError
from fpdf import FPDF
from fastapi.responses import StreamingResponse

BASE_DIR = Path(__file__).parent

app = FastAPI(title="Resumo Backend")

# Wide open for local dev — we'll lock this to the Vercel domain in step 8
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

nlp = spacy.load("en_core_web_sm")

with open(BASE_DIR / "skills_dict.json", "r", encoding="utf-8") as f:
    SURFACE_TO_CANONICAL = json.load(f)

with open(BASE_DIR / "custom_aliases.json", "r", encoding="utf-8") as f:
    CUSTOM_ALIASES = json.load(f)

# custom aliases win on conflict — they exist specifically to override/patch gaps
SURFACE_TO_CANONICAL.update(CUSTOM_ALIASES)

with open(BASE_DIR / "excluded_canonical_skills.json", "r", encoding="utf-8") as f:
    EXCLUDED_CANONICAL = set(json.load(f))

load_dotenv()
groq_client = Groq(api_key=os.environ.get("GROQ_API_KEY"))
GROQ_MODEL = "openai/gpt-oss-120b"

# Build the matcher once at startup — not per-request
skill_matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
skill_patterns = [nlp.make_doc(surface) for surface in SURFACE_TO_CANONICAL.keys()]
skill_matcher.add("SKILLS", skill_patterns)

FUZZY_MATCH_THRESHOLD = 85  # 0-100, tune later based on real test cases
MIN_LEN_FOR_FUZZY = 4  # below this, fuzzy comparisons are too noisy to trust

def clean_for_fuzzy(term: str) -> str:
    # canonical names carry disambiguating annotations like "(Software)" or
    # "(Programming Language)" that skew string-similarity comparisons — strip
    # them before comparing
    return re.sub(r"\s*\(.*?\)\s*", "", term).strip().lower()

@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/api/ocr")
async def ocr(file: UploadFile = File(...)):
    if not file.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="File must be an image")

    contents = await file.read()
    try:
        image = Image.open(io.BytesIO(contents))
        extracted_text = pytesseract.image_to_string(image)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OCR failed: {str(e)}")

    return {"extractedText": extracted_text.strip()}

def parse_pdf(contents: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(contents)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)

def parse_docx(contents: bytes) -> str:
    document = docx.Document(io.BytesIO(contents))
    parts = [p.text for p in document.paragraphs if p.text]
    # a lot of resume templates (Canva exports especially) put contact info or
    # skills inside tables, which document.paragraphs skips entirely
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells if cell.text)
    return "\n".join(parts)

def check_resume_text(filename: str, contents: bytes) -> str:
    """Parses a resume file to text, or raises the appropriate HTTPException."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        resume_text = parse_pdf(contents)
    elif lower.endswith(".docx"):
        resume_text = parse_docx(contents)
    elif lower.endswith(".doc"):
        raise HTTPException(
            status_code=400,
            detail="Legacy .doc files aren't supported — please re-save as .docx or .pdf"
        )
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    if not resume_text.strip():
        raise HTTPException(status_code=422, detail="Could not extract any text from the resume")
    return resume_text

NAME_LIKE_PATTERN = re.compile(r"^[A-Za-z][A-Za-z.'-]*(\s+[A-Za-z][A-Za-z.'-]*){0,3}$")
NOT_A_NAME = {
    "curriculum vitae", "resume", "cv", "personal details", "personal information",
    "contact information", "contact details", "resume of", "profile",
}

def extract_candidate_name(resume_text: str) -> str | None:
    """Best-effort heuristic: resumes overwhelmingly open with the candidate's
    name as the very first line. Deliberately conservative — returns None
    rather than guessing when the first non-empty line doesn't clearly look
    like a plain name (no digits, no @, no punctuation-heavy contact info,
    not a generic document-title phrase like "Curriculum Vitae"), since a
    wrong name on the report looks worse than no name at all.
    This is meant as a suggestion for the frontend to show the user for
    confirmation/editing, not a final answer — see /api/analyze's
    candidateName field and ReportRequest.candidateName."""
    for line in resume_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.lower() in NOT_A_NAME:
            continue  # skip a title line, check the next non-empty line for the actual name
        if NAME_LIKE_PATTERN.match(stripped):
            return stripped
        return None  # first substantive line doesn't look like a name — bail rather than guess
    return None

# --- Structure & completeness checks -----------------------------------

EMAIL_PATTERN = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
# find number-and-separator-heavy candidates first, then require a realistic
# digit count — avoids "2020-2023" (8 digits) being mistaken for a phone number
PHONE_CANDIDATE_PATTERN = re.compile(r"[\d][\d\s().+-]{7,}\d")

def has_phone_number(text: str) -> bool:
    for match in PHONE_CANDIDATE_PATTERN.finditer(text):
        digit_count = sum(c.isdigit() for c in match.group())
        if 9 <= digit_count <= 15:
            return True
    return False

SECTION_KEYWORDS = {
    "experience": ["experience", "employment history", "work history"],
    "education": ["education", "academic background"],
    "skills": ["skills", "technical skills", "core competencies"],
    "projects": ["projects", "personal projects", "academic projects"],
    "summary": ["summary", "objective", "profile", "about me"],
    "certifications": ["certifications", "certificates", "licenses"],
}

def detect_sections(text: str) -> set[str]:
    found = set()
    for line in text.splitlines():
        stripped = line.strip().strip(":-•* ").lower()
        if not stripped or len(stripped.split()) > 5:
            continue  # header lines are short — long lines are body text, not headers
        for category, keywords in SECTION_KEYWORDS.items():
            if any(kw in stripped for kw in keywords):
                found.add(category)
    return found

MONTH = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?"
DATE_TOKEN = rf"(?:{MONTH}\s+\d{{4}}|\d{{1,2}}/\d{{4}}|\d{{4}})"
DATE_RANGE_PATTERN = re.compile(
    rf"{DATE_TOKEN}\s*(?:-|–|—|to)\s*(?:{DATE_TOKEN}|present|current)",
    re.IGNORECASE,
)

def check_garbled_text(text: str) -> float:
    """Heuristic 0-100 — 100 means the extraction looks clean, lower means signs
    of column/table scrambling. Can't fully distinguish real scrambling from a
    legitimately terse one-item-per-line format without layout coordinates from
    pdfplumber — treat this as a soft signal, not a hard gate (hence the modest
    30% weight in structure_score below)."""
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if not lines:
        return 0.0
    non_ws = re.sub(r"\s", "", text)
    if not non_ws:
        return 0.0
    alpha_ratio = sum(c.isalpha() for c in non_ws) / len(non_ws)
    short_line_ratio = sum(1 for l in lines if len(l.split()) <= 2) / len(lines)

    alpha_score = min(alpha_ratio / 0.6, 1) * 100
    line_score = 100 if short_line_ratio <= 0.3 else max(0, 100 - (short_line_ratio - 0.3) / 0.4 * 100)
    return round(0.5 * alpha_score + 0.5 * line_score, 1)

def score_length(word_count: int) -> float:
    """Full marks in the 250-900 word range (roughly a healthy 1-2 page resume),
    tapering off outside it rather than a hard cutoff."""
    if 250 <= word_count <= 900:
        return 100.0
    if word_count < 250:
        return max(0, word_count / 250 * 100)
    return max(0, 100 - (word_count - 900) / 600 * 100)

def check_structure(text: str) -> dict:
    has_email = bool(EMAIL_PATTERN.search(text))
    has_phone = has_phone_number(text)
    contact_score = 100 if (has_email and has_phone) else (50 if (has_email or has_phone) else 0)

    sections = detect_sections(text)
    headers_score = 100 if len(sections) >= 2 else (50 if len(sections) == 1 else 0)

    garbled_score = check_garbled_text(text)

    score = round(0.4 * contact_score + 0.3 * headers_score + 0.3 * garbled_score, 1)
    return {
        "score": score,
        "hasEmail": has_email,
        "hasPhone": has_phone,
        "sectionsDetected": sorted(sections),
        "garbledTextScore": garbled_score,
    }

def check_completeness(text: str) -> dict:
    word_count = len(text.split())
    length_score = score_length(word_count)

    sections = detect_sections(text)
    slots = {
        "experience_or_projects": {"experience", "projects"},
        "education": {"education"},
        "skills": {"skills"},
    }
    slots_satisfied = sum(1 for kws in slots.values() if kws & sections)
    coverage_score = slots_satisfied / len(slots) * 100

    date_count = len(DATE_RANGE_PATTERN.findall(text))
    dated_score = min(date_count / 2, 1) * 100  # 2+ dated entries = full marks

    score = round(0.35 * length_score + 0.35 * coverage_score + 0.3 * dated_score, 1)
    return {
        "score": score,
        "wordCount": word_count,
        "sectionsDetected": sorted(sections),
        "datedEntriesFound": date_count,
    }

@app.post("/api/debug-structure")
async def debug_structure(text: str = Form(...)):
    return check_structure(text)

@app.post("/api/debug-completeness")
async def debug_completeness(text: str = Form(...)):
    return check_completeness(text)


# --- Keyword extraction & matching --------------------------------------

def extract_keywords(text: str) -> dict[str, int]:
    # make_doc only tokenizes — same as what we already use to build the matcher's
    # patterns — so we skip the tagger/parser/NER, which PhraseMatcher never needed anyway
    doc = nlp.make_doc(text)
    matches = skill_matcher(doc)

    # PhraseMatcher returns overlapping/nested matches by default (e.g. "full stack"
    # AND "full stack developer" both hitting on the same words) — filter_spans keeps
    # only the longest non-overlapping span so we don't double-count
    spans = filter_spans([doc[start:end] for _, start, end in matches])

    counts: dict[str, int] = {}
    for span in spans:
        surface_form = span.text.lower()
        canonical = SURFACE_TO_CANONICAL.get(surface_form, surface_form)
         # Ignore skills to not score
        if canonical in EXCLUDED_CANONICAL:
            continue
        counts[canonical] = counts.get(canonical, 0) + 1

    return counts

@app.post("/api/debug-keywords")
async def debug_keywords(text: str = Form(...)):
    return extract_keywords(text)

def match_keywords(jd_keywords: dict[str, int], resume_keywords: dict[str, int]):
    jd_terms = set(jd_keywords)
    resume_terms = set(resume_keywords)

    exact_matched = jd_terms & resume_terms
    remaining_jd = jd_terms - exact_matched

    fuzzy_matched = set()
    for jd_term in remaining_jd:
        jd_clean = clean_for_fuzzy(jd_term)
        if len(jd_clean) < MIN_LEN_FOR_FUZZY:
            continue
        for resume_term in resume_terms:
            resume_clean = clean_for_fuzzy(resume_term)
            if len(resume_clean) < MIN_LEN_FOR_FUZZY:
                continue
            # ratio, not partial_ratio: partial_ratio only needs the shorter string
            # to align against SOME substring of the longer one, which scores things
            # like java/javascript, react/reactor, and go/django all at 100 — exactly
            # the skill confusions this is supposed to avoid
            if fuzz.ratio(jd_clean, resume_clean) >= FUZZY_MATCH_THRESHOLD:
                fuzzy_matched.add(jd_term)
                break

    matched = exact_matched | fuzzy_matched
    missing = jd_terms - matched

    # weighted overlap: JD terms mentioned more often count for more
    total_jd_weight = sum(jd_keywords.values())
    matched_weight = sum(jd_keywords[term] for term in matched)
    keyword_score = (matched_weight / total_jd_weight * 100) if total_jd_weight > 0 else 0

    return {
        "matchedKeywords": sorted(matched),
        "missingKeywords": sorted(missing),
        "exactMatchCount": len(exact_matched),
        "fuzzyMatchCount": len(fuzzy_matched),
        "keywordScore": round(keyword_score, 1),
    }

@app.post("/api/debug-match")
async def debug_match(jdText: str = Form(...), resumeText: str = Form(...)):
    jd_keywords = extract_keywords(jdText)
    resume_keywords = extract_keywords(resumeText)
    return match_keywords(jd_keywords, resume_keywords)


# --- Full pipeline -------------------------------------------------------

KEYWORD_WEIGHT = 0.5
STRUCTURE_WEIGHT = 0.3
COMPLETENESS_WEIGHT = 0.2

@app.post("/api/analyze")
async def analyze(resume: UploadFile = File(...), jdText: str = Form(...)):
    contents = await resume.read()
    resume_text = check_resume_text(resume.filename or "", contents)

    jd_keywords = extract_keywords(jdText)
    resume_keywords = extract_keywords(resume_text)
    keyword_result = match_keywords(jd_keywords, resume_keywords)

    structure_result = check_structure(resume_text)
    completeness_result = check_completeness(resume_text)

    final_score = (
        KEYWORD_WEIGHT * keyword_result["keywordScore"]
        + STRUCTURE_WEIGHT * structure_result["score"]
        + COMPLETENESS_WEIGHT * completeness_result["score"]
    )

    return {
        "resumeText": resume_text,
        "candidateName": extract_candidate_name(resume_text),
        "finalScore": round(final_score, 1),
        "breakdown": {
            "keywordScore": keyword_result["keywordScore"],
            "structureScore": structure_result["score"],
            "completenessScore": completeness_result["score"],
        },
        "matchedKeywords": keyword_result["matchedKeywords"],
        "missingKeywords": keyword_result["missingKeywords"],
        "structureDetails": structure_result,
        "completenessDetails": completeness_result,
    }

class Breakdown(BaseModel):
    keywordScore: float
    structureScore: float
    completenessScore: float

class StructureDetails(BaseModel):
    score: float
    hasEmail: bool
    hasPhone: bool
    sectionsDetected: list[str]
    garbledTextScore: float

class CompletenessDetails(BaseModel):
    score: float
    wordCount: int
    sectionsDetected: list[str]
    datedEntriesFound: int

class Suggestions(BaseModel):
    summarySuggestion: str
    bulletSuggestions: list[str]
    generalTips: list[str]

class SuggestionsRequest(BaseModel):
    resumeText: str
    jdText: str
    missingKeywords: list[str]
    keywordScore: float
    structureDetails: StructureDetails
    completenessDetails: CompletenessDetails

SUGGESTIONS_SYSTEM_PROMPT = """You are a resume optimization assistant. Given a resume, a job \
description, and a list of keywords missing from the resume, produce specific, honest suggestions.

CRITICAL RULE: When rewriting a bullet point, you may ONLY rephrase, restructure, or better-word what \
is ALREADY stated in that bullet's source text. You must NEVER add a technology, tool, skill, outcome, \
or metric (including percentages or numbers) that does not appear, explicitly or very directly implied, \
in the candidate's original resume text. This applies even if that keyword is in the missing-keywords \
list — inserting an unearned keyword into a rewritten bullet is a serious error, not a helpful suggestion.

If a missing keyword genuinely isn't supported by anything in the resume, do NOT put it in a bullet — \
only mention it, conditionally, in generalTips (e.g. "if you have experience with Docker, consider \
adding it"), exactly as you already do correctly for generalTips.

Respond ONLY with valid JSON matching this exact schema, no markdown fences, no preamble:

{
  "summarySuggestion": "a 2-3 sentence tailored professional summary based only on the candidate's actual experience",
  "bulletSuggestions": ["3-5 rewritten bullet points — rephrasing ONLY, no invented technologies, tools, or metrics"],
  "generalTips": ["2-4 short, actionable tips, phrased conditionally for any skill not evidenced in the resume"]
}"""

@app.post("/api/suggestions")
async def suggestions(req: SuggestionsRequest):
    structure_summary = (
        f"Email detected: {req.structureDetails.hasEmail}. "
        f"Phone detected: {req.structureDetails.hasPhone}. "
        f"Sections detected: {', '.join(req.structureDetails.sectionsDetected) or 'none'}."
    )
    completeness_summary = (
        f"Word count: {req.completenessDetails.wordCount}. "
        f"Dated experience entries found: {req.completenessDetails.datedEntriesFound}."
    )
    user_prompt = f"""RESUME:
{req.resumeText}

JOB DESCRIPTION:
{req.jdText}

MISSING KEYWORDS: {", ".join(req.missingKeywords) if req.missingKeywords else "none"}
CURRENT KEYWORD MATCH SCORE: {req.keywordScore}%

RESUME STRUCTURE FINDINGS: {structure_summary}
RESUME COMPLETENESS FINDINGS: {completeness_summary}
(generalTips may reference the structure/completeness findings above — e.g. a
missing phone number or thin word count — not just missing keywords.)"""

    last_error = None
    for attempt in range(2):  # one retry if the model returns a malformed shape
        try:
            completion = groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {"role": "system", "content": SUGGESTIONS_SYSTEM_PROMPT},
                    {"role": "user", "content": user_prompt},
                ],
                response_format={"type": "json_object"},
                max_completion_tokens=1500,
                temperature=0.4,
            )
            raw = completion.choices[0].message.content
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"Suggestion generation failed: {str(e)}")

        try:
            parsed = json.loads(raw)
            # json_object mode only guarantees valid JSON syntax, not that it
            # matches our schema — validate before trusting it, since this same
            # payload gets passed straight into ReportRequest.suggestions later
            return Suggestions.model_validate(parsed).model_dump()
        except (json.JSONDecodeError, ValidationError) as e:
            last_error = e
            continue

    raise HTTPException(status_code=502, detail=f"Model response didn't match the expected format after retry: {last_error}")

# --- PDF report generation ------------------------------------------------

SMART_CHAR_MAP = {
    "\u2018": "'", "\u2019": "'",
    "\u201c": '"', "\u201d": '"',
    "\u2013": "-", "\u2014": "-", "\u2011": "-", "\u2212": "-",
    "\u2026": "...",
    "\u2022": "-", "\u25cf": "-", "\u2605": "*",
}

def sanitize_for_pdf(text: str) -> str:
    # fpdf2's default core fonts only support Latin-1 — LLM output routinely
    # contains "smart" typographic characters (curly quotes, en/em dashes,
    # non-breaking hyphens) that crash multi_cell() otherwise. Replace the
    # common ones with sane equivalents, then use encode/replace as a
    # last-resort safety net for anything unanticipated.
    for bad, good in SMART_CHAR_MAP.items():
        text = text.replace(bad, good)
    return text.encode("latin-1", errors="replace").decode("latin-1")

class ReportRequest(BaseModel):
    finalScore: float
    breakdown: Breakdown
    matchedKeywords: list[str]
    missingKeywords: list[str]
    structureDetails: StructureDetails
    completenessDetails: CompletenessDetails
    suggestions: Suggestions | None = None  # optional — report should still work if /api/suggestions failed/was skipped
    candidateName: str | None = None

class ResumoReportPDF(FPDF):
    def footer(self):
        self.set_y(-15)
        self.set_draw_color(210, 210, 210)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(2)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(130, 130, 130)
        self.cell(0, 8, f"Generated by Resumo, a project of Muhammad Usman  |  Page {self.page_no()}", align="C")
        self.set_text_color(0, 0, 0)

def score_color(score: float) -> tuple[int, int, int]:
    if score >= 75:
        return (46, 139, 87)   # green
    if score >= 50:
        return (200, 150, 20)  # amber
    return (196, 60, 60)       # red

def divider(pdf: FPDF, gap_before: float = 2, gap_after: float = 4):
    pdf.ln(gap_before)
    pdf.set_draw_color(210, 210, 210)
    y = pdf.get_y()
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.ln(gap_after)

def overall_score_box(pdf: FPDF, score: float):
    # rect() never touches the text cursor, so we track x/y ourselves throughout
    # rather than relying on any auto-advance behavior
    r, g, b = score_color(score)
    box_x, box_y = pdf.get_x(), pdf.get_y()
    box_w, box_h = 65, 26
    pdf.set_fill_color(r, g, b)
    pdf.rect(box_x, box_y, box_w, box_h, style="F")
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_xy(box_x, box_y + 4)
    pdf.cell(box_w, 12, f"{score:.1f}/100", align="C")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_xy(box_x, box_y + 17)
    pdf.cell(box_w, 6, "Overall Match Score", align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.set_xy(pdf.l_margin, box_y + box_h + 6)

def draw_score_bar(pdf: FPDF, label: str, score: float, bar_width: float = 170, bar_height: float = 7):
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 6, f"{label}: {score:.1f}/100", new_x="LMARGIN", new_y="NEXT")
    x, y = pdf.get_x(), pdf.get_y()
    pdf.set_draw_color(190, 190, 190)
    pdf.set_fill_color(235, 235, 235)
    pdf.rect(x, y, bar_width, bar_height, style="DF")
    filled = bar_width * max(0, min(score, 100)) / 100
    if filled > 0:
        r, g, b = score_color(score)
        pdf.set_fill_color(r, g, b)
        pdf.rect(x, y, filled, bar_height, style="F")
    pdf.set_xy(x, y + bar_height + 5)

def keyword_box(pdf: FPDF, keywords: list[str], border_color: tuple, fill_color: tuple, empty_text: str):
    pdf.set_draw_color(*border_color)
    pdf.set_fill_color(*fill_color)
    pdf.set_font("Helvetica", "", 10)
    text = ", ".join(keywords) if keywords else empty_text
    pdf.multi_cell(0, 6, sanitize_for_pdf(text), border=1, fill=True, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

def build_report_pdf(data: ReportRequest) -> bytes:
    pdf = ResumoReportPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # NOTE: new_x/new_y must be passed explicitly on every multi_cell call —
    # without it, the cursor doesn't reliably reset to the left margin between
    # calls in this fpdf2 version, and consecutive bullets drift toward the
    # page edge until there's no room left to render a character at all.
    # rect()-based elements (score box/bars) instead track x/y explicitly via
    # get_x()/get_y()/set_xy(), since rect() never touches the cursor at all.
    def h1(text):
        pdf.set_font("Helvetica", "B", 18)
        pdf.multi_cell(0, 10, sanitize_for_pdf(text), new_x="LMARGIN", new_y="NEXT")

    def h2(text):
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(0, 8, sanitize_for_pdf(text), new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)

    def body(text):
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 6, sanitize_for_pdf(text), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

    def bullet(text):
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 6, sanitize_for_pdf(f"- {text}"), new_x="LMARGIN", new_y="NEXT")

    title = f"Resumo Analysis Report for {data.candidateName}" if data.candidateName else "Resumo Analysis Report"
    h1(title)
    pdf.ln(2)

    overall_score_box(pdf, data.finalScore)
    divider(pdf)

    h2("Score Breakdown")
    draw_score_bar(pdf, "Keyword Match (50%)", data.breakdown.keywordScore)
    draw_score_bar(pdf, "Resume Structure (30%)", data.breakdown.structureScore)
    draw_score_bar(pdf, "Completeness (20%)", data.breakdown.completenessScore)
    divider(pdf)

    h2("Matched Keywords")
    keyword_box(pdf, data.matchedKeywords, border_color=(120, 180, 120),
                fill_color=(232, 245, 232), empty_text="None found.")

    h2("Missing Keywords")
    keyword_box(pdf, data.missingKeywords, border_color=(210, 140, 140),
                fill_color=(250, 232, 232), empty_text="None - great coverage!")

    if data.suggestions:
        divider(pdf)
        h2("Suggested Summary")
        body(data.suggestions.summarySuggestion)
        h2("Suggested Bullet Points")
        for b in data.suggestions.bulletSuggestions:
            bullet(b)
        h2("General Tips")
        for t in data.suggestions.generalTips:
            bullet(t)

    divider(pdf)
    h2("Formatting Checklist")
    bullet(f"{'[x]' if data.structureDetails.hasEmail else '[ ]'} Email address detected")
    bullet(f"{'[x]' if data.structureDetails.hasPhone else '[ ]'} Phone number detected")
    bullet(f"Sections detected: {', '.join(data.structureDetails.sectionsDetected) or 'none'}")
    bullet(f"Word count: {data.completenessDetails.wordCount}")
    bullet(f"Dated entries found: {data.completenessDetails.datedEntriesFound}")

    return bytes(pdf.output())

@app.post("/api/report")
async def report(req: ReportRequest):
    pdf_bytes = build_report_pdf(req)
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=resumo_report.pdf"},
    )